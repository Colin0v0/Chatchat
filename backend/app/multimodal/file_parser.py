from __future__ import annotations

import csv
import json
from pathlib import Path

from ..storage.media import MEDIA_ROOT
from ..storage.models import MessageAttachment
from .file_types import is_text_attachment


class FileParser:
    def __init__(self, *, text_max_chars: int, table_row_limit: int, table_column_limit: int):
        self._text_max_chars = max(1, text_max_chars)
        self._table_row_limit = max(1, table_row_limit)
        self._table_column_limit = max(1, table_column_limit)

    def extract_markdown(self, attachments: list[MessageAttachment]) -> str:
        blocks = [
            self._build_attachment_block(index=index, attachment=attachment)
            for index, attachment in enumerate(attachments, start=1)
            if attachment.kind == "file"
        ]
        return "\n\n".join(blocks).strip()

    def _build_attachment_block(self, *, index: int, attachment: MessageAttachment) -> str:
        file_path = MEDIA_ROOT / Path(attachment.relative_path)
        content = self._parse_file(attachment=attachment, file_path=file_path)
        return "\n".join(
            [
                f"### File {index}",
                f"name: {attachment.original_name}",
                f"type: {attachment.extension or Path(attachment.original_name).suffix.lower() or attachment.mime_type}",
                "content:",
                self._truncate(content),
            ]
        ).strip()

    def _parse_file(self, *, attachment: MessageAttachment, file_path: Path) -> str:
        extension = (attachment.extension or file_path.suffix).lower()
        if extension == ".pdf":
            return self._parse_pdf(file_path)
        if extension == ".csv":
            return self._parse_csv(file_path)
        if extension == ".xlsx":
            return self._parse_xlsx(file_path)
        if extension == ".docx":
            return self._parse_docx(file_path)
        if is_text_attachment(extension):
            return self._parse_text(file_path, extension)
        raise RuntimeError(f"Unsupported file parser for {attachment.original_name}.")

    def _parse_pdf(self, file_path: Path) -> str:
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise RuntimeError("PDF parsing requires pypdf in the backend environment.") from exc

        reader = PdfReader(str(file_path))
        blocks: list[str] = []
        for index, page in enumerate(reader.pages, start=1):
            text = " ".join((page.extract_text() or "").split()).strip()
            if not text:
                continue
            blocks.append(f"#### Page {index}\n{text}")
        if not blocks:
            raise RuntimeError(f"No readable text found in PDF: {file_path.name}.")
        return "\n\n".join(blocks)

    def _parse_text(self, file_path: Path, extension: str) -> str:
        raw = file_path.read_bytes()
        decoded = raw.decode("utf-8", errors="replace").strip()
        if not decoded:
            raise RuntimeError(f"No readable text found in {file_path.name}.")

        if extension == ".json":
            try:
                parsed = json.loads(decoded)
            except json.JSONDecodeError:
                return decoded
            return json.dumps(parsed, ensure_ascii=False, indent=2)

        return decoded

    def _parse_csv(self, file_path: Path) -> str:
        with file_path.open("r", encoding="utf-8", errors="replace", newline="") as handle:
            reader = csv.reader(handle)
            rows = [row[: self._table_column_limit] for row in reader]

        return self._format_rows(rows=rows, label=file_path.name)

    def _parse_xlsx(self, file_path: Path) -> str:
        try:
            from openpyxl import load_workbook
        except ImportError as exc:
            raise RuntimeError("XLSX parsing requires openpyxl in the backend environment.") from exc

        workbook = load_workbook(filename=str(file_path), read_only=True, data_only=True)
        sheet_blocks: list[str] = []
        for sheet in workbook.worksheets:
            rows: list[list[str]] = []
            for row in sheet.iter_rows(values_only=True):
                values = ["" if value is None else str(value) for value in row[: self._table_column_limit]]
                rows.append(values)
                if len(rows) >= self._table_row_limit:
                    break
            if rows:
                sheet_blocks.append(self._format_rows(rows=rows, label=f"{file_path.name} / {sheet.title}"))

        if not sheet_blocks:
            raise RuntimeError(f"No readable worksheet data found in {file_path.name}.")
        return "\n\n".join(sheet_blocks)

    def _parse_docx(self, file_path: Path) -> str:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("DOCX parsing requires python-docx in the backend environment.") from exc

        document = Document(str(file_path))
        blocks: list[str] = []
        for paragraph in document.paragraphs:
            text = " ".join(paragraph.text.split()).strip()
            if text:
                blocks.append(text)

        for table_index, table in enumerate(document.tables, start=1):
            rows: list[list[str]] = []
            for row in table.rows[: self._table_row_limit]:
                rows.append(
                    [
                        " ".join(cell.text.split()).strip()
                        for cell in row.cells[: self._table_column_limit]
                    ]
                )
            if rows:
                blocks.append(f"#### Table {table_index}\n{self._format_rows(rows=rows, label='table')}")

        if not blocks:
            raise RuntimeError(f"No readable content found in {file_path.name}.")
        return "\n\n".join(blocks)

    def _format_rows(self, *, rows: list[list[str]], label: str) -> str:
        if not rows:
            raise RuntimeError(f"No readable tabular content found in {label}.")

        header = rows[0]
        body = rows[1:self._table_row_limit]
        rendered_rows = [
            f"- columns: {', '.join(item or '<empty>' for item in header)}",
        ]
        for index, row in enumerate(body, start=1):
            rendered_rows.append(f"- row {index}: {', '.join(item or '<empty>' for item in row)}")
        return "\n".join(rendered_rows)

    def _truncate(self, content: str) -> str:
        normalized = content.strip()
        if len(normalized) <= self._text_max_chars:
            return normalized
        cutoff = max(0, self._text_max_chars - 3)
        return normalized[:cutoff].rstrip() + "..."
